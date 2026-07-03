import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

QR_DIR = os.path.join(os.path.dirname(__file__), "static", "qrs")
OUTPUT = os.path.join(os.path.dirname(__file__), "codigos_qr_habitaciones.docx")


def numero_habitacion(filename):
    m = re.search(r"habitacion_(\d+)\.png", filename)
    return int(m.group(1)) if m else 0


def main():
    archivos = sorted(
        [f for f in os.listdir(QR_DIR) if f.endswith(".png")],
        key=numero_habitacion,
    )

    doc = Document()

    # Márgenes pequeños para que el QR se vea grande
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    for i, archivo in enumerate(archivos):
        num = numero_habitacion(archivo)
        ruta = os.path.join(QR_DIR, archivo)

        if i > 0:
            doc.add_page_break()

        # Párrafo con el QR centrado
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(ruta, width=Inches(4.5))

        # Párrafo con el número de habitación
        p_txt = doc.add_paragraph()
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_txt = p_txt.add_run(f"Habitación {num}")
        run_txt.font.size = Pt(28)
        run_txt.font.bold = True

    doc.save(OUTPUT)
    print(f"Documento guardado: {OUTPUT}")
    print(f"Total de habitaciones: {len(archivos)}")


if __name__ == "__main__":
    main()
